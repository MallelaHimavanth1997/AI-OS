"""Tailor parsed resumes to target job descriptions without inventing experience."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Sequence
import re

from agents.job_matching_agent import MatchBreakdown
from agents.job_search_agent import ParsedJob
from agents.resume_intelligence_agent import ParsedResume


@dataclass(slots=True)
class TailoredResumeDocument:
    """Structured ATS-safe resume output."""

    text: str
    ats_score: float
    matched_keywords: list[str] = field(default_factory=list)
    missing_keywords: list[str] = field(default_factory=list)
    recommendations: list[str] = field(default_factory=list)
    title: str = "Tailored Resume"
    candidate_name: str = ""
    source_job_title: str = ""


class ResumeTailor:
    """Create tailored resume content from parsed resume and job inputs."""

    def build(self, resume: ParsedResume, job: ParsedJob, match: MatchBreakdown) -> TailoredResumeDocument:
        job_keywords = self._normalize_keywords(job.keywords + self._extract_keywords(job.description) + self._extract_keywords(job.raw_text))
        resume_skill_map = self._build_reserve_lookup(resume)
        prioritized_keywords = self._prioritize_keywords(job_keywords, resume_skill_map, max_keywords=15)
        matched_keywords = [keyword for keyword in prioritized_keywords if self._contains_keyword(resume.raw_text, keyword) or keyword in resume_skill_map]
        missing_keywords = [keyword for keyword in prioritized_keywords if keyword not in matched_keywords]

        summary = self._build_summary(resume, prioritized_keywords, match)
        skills_section = self._build_skills_section(resume, prioritized_keywords)
        experience_section = self._build_experience_section(resume, prioritized_keywords)
        projects_section = self._build_projects_section(resume, prioritized_keywords)
        certifications_section = self._build_certifications_section(resume)
        education_section = self._build_education_section(resume)

        ats_score = self._compute_ats_score(matched_keywords, prioritized_keywords, match)
        text = self._render_document(
            resume=resume,
            job=job,
            summary=summary,
            skills_section=skills_section,
            experience_section=experience_section,
            projects_section=projects_section,
            certifications_section=certifications_section,
            education_section=education_section,
        )

        recommendations = self._build_recommendations(missing_keywords, resume, job)

        return TailoredResumeDocument(
            text=text,
            ats_score=ats_score,
            matched_keywords=matched_keywords,
            missing_keywords=missing_keywords,
            recommendations=recommendations,
            title=f"Tailored Resume - {job.title}",
            candidate_name=resume.full_name,
            source_job_title=job.title,
        )

    def _build_summary(self, resume: ParsedResume, keywords: list[str], match: MatchBreakdown) -> str:
        base_summary = resume.summary.strip() if resume.summary else ""
        focus_keywords = [keyword for keyword in keywords[:5] if keyword]
        keyword_phrase = ", ".join(focus_keywords)

        summary_parts: list[str] = []
        if base_summary:
            summary_parts.append(base_summary)
        if keyword_phrase:
            summary_parts.append(f"Focused on {keyword_phrase}.")
        summary_parts.append(f"Target fit score: {match.match_score:.2f}%.")
        return " ".join(summary_parts).strip()

    def _build_skills_section(self, resume: ParsedResume, keywords: list[str]) -> list[str]:
        ordered = self._merge_preserving_order([*keywords, *resume.skills])
        return [skill for skill in ordered[:20]]

    def _build_experience_section(self, resume: ParsedResume, keywords: list[str]) -> list[str]:
        return self._select_bullets(resume.experience, keywords, limit=10)

    def _build_projects_section(self, resume: ParsedResume, keywords: list[str]) -> list[str]:
        return self._select_bullets(resume.projects, keywords, limit=8)

    def _build_certifications_section(self, resume: ParsedResume) -> list[str]:
        return self._select_bullets(resume.certifications, [], limit=5)

    def _build_education_section(self, resume: ParsedResume) -> list[str]:
        return self._select_bullets(resume.education, [], limit=5)

    def _select_bullets(self, items: Sequence[str], keywords: list[str], *, limit: int) -> list[str]:
        scored = sorted(items, key=lambda item: self._score_text(item, keywords), reverse=True)
        return [self._emphasize_keywords(item, keywords) for item in scored[:limit]]

    def _render_document(
        self,
        *,
        resume: ParsedResume,
        job: ParsedJob,
        summary: str,
        skills_section: list[str],
        experience_section: list[str],
        projects_section: list[str],
        certifications_section: list[str],
        education_section: list[str],
    ) -> str:
        sections: list[str] = [
            resume.full_name,
            resume.headline or job.title,
            job.company,
            f"Target Role: {job.title}",
            f"Location: {job.location}",
            "",
            "SUMMARY",
            summary,
            "",
            "SKILLS",
            self._render_bullets(skills_section),
            "",
            "EXPERIENCE",
            self._render_bullets(experience_section),
        ]
        if projects_section:
            sections.extend(["", "PROJECTS", self._render_bullets(projects_section)])
        if certifications_section:
            sections.extend(["", "CERTIFICATIONS", self._render_bullets(certifications_section)])
        if education_section:
            sections.extend(["", "EDUCATION", self._render_bullets(education_section)])
        return "\n".join(section for section in sections if section is not None).strip()

    def _render_bullets(self, items: Sequence[str]) -> str:
        return "\n".join(f"- {item}" for item in items)

    def _build_recommendations(self, missing_keywords: list[str], resume: ParsedResume, job: ParsedJob) -> list[str]:
        recommendations: list[str] = []
        if missing_keywords:
            recommendations.append(f"Emphasize these existing keywords where truthful: {', '.join(missing_keywords[:8])}.")
        if not resume.summary:
            recommendations.append("Add a concise professional summary for faster ATS parsing.")
        if not resume.skills:
            recommendations.append("Keep a dedicated skills section near the top of the resume.")
        if job.title and job.title.lower() not in resume.headline.lower():
            recommendations.append("Align the headline more closely with the target role.")
        return recommendations

    def _compute_ats_score(self, matched_keywords: list[str], prioritized_keywords: list[str], match: MatchBreakdown) -> float:
        if not prioritized_keywords:
            return round(match.match_score, 2)
        keyword_coverage = len(matched_keywords) / len(prioritized_keywords)
        score = (match.match_score * 0.7) + (keyword_coverage * 100.0 * 0.3)
        return round(min(score, 100.0), 2)

    def _prioritize_keywords(self, job_keywords: list[str], resume_skill_map: set[str], *, max_keywords: int) -> list[str]:
        prioritized: list[str] = []
        for keyword in job_keywords:
            if keyword in resume_skill_map and keyword not in prioritized:
                prioritized.append(keyword)
        for keyword in job_keywords:
            if keyword not in prioritized:
                prioritized.append(keyword)
        return prioritized[:max_keywords]

    def _build_reserve_lookup(self, resume: ParsedResume) -> set[str]:
        values = resume.skills + resume.keywords + resume.certifications + resume.projects + resume.experience + resume.education
        return {self._normalize_text(value) for value in values if value}

    def _extract_keywords(self, text: str) -> list[str]:
        tokens = re.findall(r"[A-Za-z][A-Za-z0-9+#\.-]{1,}", text.lower())
        stop_words = {"and", "the", "for", "with", "from", "that", "this", "you", "your", "our", "are", "will", "have", "into", "over", "role", "job", "team", "work"}
        return self._normalize_keywords([token for token in tokens if token not in stop_words])

    def _normalize_keywords(self, keywords: Sequence[str]) -> list[str]:
        result: list[str] = []
        seen: set[str] = set()
        for keyword in keywords:
            normalized = self._normalize_text(keyword)
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            result.append(normalized)
        return result

    def _merge_preserving_order(self, values: Sequence[str]) -> list[str]:
        result: list[str] = []
        seen: set[str] = set()
        for value in values:
            normalized = self._normalize_text(value)
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            result.append(value)
        return result

    def _score_text(self, value: str, keywords: list[str]) -> int:
        normalized_value = self._normalize_text(value)
        return sum(1 for keyword in keywords if keyword in normalized_value)

    def _emphasize_keywords(self, value: str, keywords: list[str]) -> str:
        result = value
        for keyword in sorted(keywords, key=len, reverse=True):
            if not keyword:
                continue
            pattern = re.compile(re.escape(keyword), re.IGNORECASE)
            result = pattern.sub(lambda match: match.group(0).upper(), result)
        return result

    def _contains_keyword(self, text: str, keyword: str) -> bool:
        return self._normalize_text(keyword) in self._normalize_text(text)

    def _normalize_text(self, value: str) -> str:
        return re.sub(r"\s+", " ", value.strip().lower())

    def export_docx(self, document: TailoredResumeDocument, output_path: str | Path) -> Path:
        """Export the tailored resume to DOCX format."""

        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - import check
            raise RuntimeError("python-docx is required to export DOCX resumes.") from exc

        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading(document.candidate_name or document.title, level=0)
        if document.source_job_title:
            doc.add_paragraph(f"Target Role: {document.source_job_title}")
        doc.add_paragraph(f"ATS Score: {document.ats_score:.2f}%")

        current_heading = None
        for line in document.text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped in {"SUMMARY", "SKILLS", "EXPERIENCE", "PROJECTS", "CERTIFICATIONS", "EDUCATION"}:
                current_heading = stripped
                doc.add_heading(stripped.title(), level=1)
                continue
            if stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif current_heading is None:
                doc.add_paragraph(stripped)
            else:
                doc.add_paragraph(stripped)

        if document.recommendations:
            doc.add_heading("Recommendations", level=1)
            for recommendation in document.recommendations:
                doc.add_paragraph(recommendation, style="List Bullet")

        doc.save(str(path))
        return path

    def export_pdf(self, document: TailoredResumeDocument, output_path: str | Path) -> Path:
        """Export the tailored resume to PDF format."""

        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
        except ImportError as exc:  # pragma: no cover - import check
            raise RuntimeError("reportlab is required to export PDF resumes.") from exc

        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        pdf = canvas.Canvas(str(path), pagesize=letter)
        width, height = letter
        margin = 48
        y = height - margin

        def write_line(text: str, *, bold: bool = False) -> None:
            nonlocal y
            if y < margin:
                pdf.showPage()
                y = height - margin
            pdf.setFont("Helvetica-Bold" if bold else "Helvetica", 11 if bold else 10)
            pdf.drawString(margin, y, text[:110])
            y -= 14

        write_line(document.candidate_name or document.title, bold=True)
        if document.source_job_title:
            write_line(f"Target Role: {document.source_job_title}")
        write_line(f"ATS Score: {document.ats_score:.2f}%")
        y -= 4

        for line in document.text.splitlines():
            stripped = line.strip()
            if not stripped:
                y -= 4
                continue
            write_line(stripped, bold=stripped in {"SUMMARY", "SKILLS", "EXPERIENCE", "PROJECTS", "CERTIFICATIONS", "EDUCATION"})

        if document.recommendations:
            y -= 8
            write_line("Recommendations", bold=True)
            for recommendation in document.recommendations:
                write_line(f"- {recommendation}")

        pdf.save()
        return path